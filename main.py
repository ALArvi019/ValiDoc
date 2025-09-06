import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import csv
import re
import os
import datetime
import threading
import platform
import subprocess
import sys

# pip install python-docx docx2pdf endesive cryptography requests asn1crypto
# On Windows pip install pywin32

from docx import Document
from docx2pdf import convert
from endesive import pdf
from cryptography.hazmat.primitives.serialization import pkcs12
from cryptography.hazmat.backends import default_backend

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from cryptography import x509

class PDFGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ValiDoc 1.0")
        self.root.geometry("700x550")
        
        # Variables
        self.doc_path = tk.StringVar()
        self.csv_path = tk.StringVar()
        self.cert_path = tk.StringVar()
        self.cert_password = tk.StringVar()
        self.sign_document = tk.BooleanVar(value=True)
        
        self.create_widgets()
        
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        title_label = ttk.Label(main_frame, text="Generador de PDFs con Firma", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        ttk.Label(main_frame, text="Plantilla Word (.docx):").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.doc_path, width=50).grid(row=1, column=1, sticky=(tk.W, tk.E), padx=5)
        ttk.Button(main_frame, text="Examinar", command=self.browse_doc).grid(row=1, column=2)
        
        ttk.Label(main_frame, text="Datos (.csv):").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.csv_path, width=50).grid(row=2, column=1, sticky=(tk.W, tk.E), padx=5)
        ttk.Button(main_frame, text="Examinar", command=self.browse_csv).grid(row=2, column=2)
        
        ttk.Label(main_frame, text="Certificado (.pfx, .p12):").grid(row=3, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.cert_path, width=50).grid(row=3, column=1, sticky=(tk.W, tk.E), padx=5)
        ttk.Button(main_frame, text="Examinar", command=self.browse_cert).grid(row=3, column=2)
        
        ttk.Label(main_frame, text="Contraseña certificado:").grid(row=4, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.cert_password, show="*", width=50).grid(row=4, column=1, sticky=(tk.W, tk.E), padx=5)
        
        self.sign_checkbox = ttk.Checkbutton(main_frame, text="Firmar documento con certificado", variable=self.sign_document, onvalue=True, offvalue=False, command=self.toggle_cert_fields)
        self.sign_checkbox.grid(row=5, column=0, columnspan=3, pady=10, sticky=tk.W)

        self.process_btn = ttk.Button(main_frame, text="Generar PDFs", command=self.start_processing_thread)
        self.process_btn.grid(row=6, column=0, columnspan=3, pady=20)
        
        self.progress = ttk.Progressbar(main_frame, orient='horizontal', mode='determinate')
        self.progress.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        log_frame = ttk.LabelFrame(main_frame, text="Log de Operación", padding=10)
        log_frame.grid(row=8, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.rowconfigure(8, weight=1)
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        
        self.log_text = tk.Text(log_frame, height=10, width=80, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        self.toggle_cert_fields()
            
    def browse_doc(self):
        filename = filedialog.askopenfilename(title="Seleccionar plantilla Word", filetypes=[("Word files", "*.docx")])
        if filename: self.doc_path.set(filename)
            
    def browse_csv(self):
        filename = filedialog.askopenfilename(title="Seleccionar archivo CSV", filetypes=[("CSV files", "*.csv")])
        if filename: self.csv_path.set(filename)
            
    def browse_cert(self):
        filename = filedialog.askopenfilename(title="Seleccionar certificado", filetypes=[("PFX/P12 files", "*.pfx *.p12")])
        if filename: self.cert_path.set(filename)

    def toggle_cert_fields(self):
        """Activa o desactiva los campos de certificado según el checkbox."""
        state = 'normal' if self.sign_document.get() else 'disabled'
        main_frame = self.root.winfo_children()[0]
        main_frame.grid_slaves(row=3, column=1)[0].config(state=state)
        main_frame.grid_slaves(row=3, column=2)[0].config(state=state)
        main_frame.grid_slaves(row=4, column=1)[0].config(state=state)

    def log(self, message):
        self.root.after(0, self._log_thread_safe, message)
        
    def _log_thread_safe(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        
    def update_progress(self, value):
        self.root.after(0, self.progress.config, {'value': value})

    def start_processing_thread(self):
        """Inicia el procesamiento en un hilo separado para no bloquear la GUI."""
        thread = threading.Thread(target=self.process_files, daemon=True)
        thread.start()
        
    def process_files(self):
        if not all([self.doc_path.get(), self.csv_path.get()]):
            self.root.after(0, lambda: messagebox.showerror("Error", "Debe seleccionar la plantilla Word y el archivo CSV."))
            return
            
        if self.sign_document.get() and not self.cert_path.get():
            self.root.after(0, lambda: messagebox.showerror("Error", "La firma está activada, pero no ha seleccionado un certificado."))
            return
            
        try:
            self.process_btn.config(state='disabled')
            self.progress['value'] = 0
            self.log("▶️ Iniciando procesamiento...")
            
            campos = self.extraer_campos_doc(self.doc_path.get())
            self.log(f"Campos encontrados: {', '.join(campos)}")
            
            datos_csv = self.validar_csv(self.csv_path.get(), campos)
            total_filas = len(datos_csv)
            self.progress['maximum'] = total_filas

            header_columna_a = list(datos_csv[0].keys())[0]
            self.log(f"Se usará la columna '{header_columna_a}' para nombrar los archivos PDF.")
            
            output_dir = "pdfs_generados"
            os.makedirs(output_dir, exist_ok=True)
            
            private_key, cert = (None, None)
            if self.sign_document.get():
                self.log("Cargando certificado digital...")
                private_key, cert = self.load_certificate(self.cert_path.get(), self.cert_password.get())
            
            for i, fila in enumerate(datos_csv):
                self.log(f"\n--- Procesando fila {i+1} de {total_filas} ---")
                
                # 1. Construir el nombre base a partir del CSV y limpiarlo
                valor_columna_a = fila.get(header_columna_a, f'fila_{i+1}')
                nombre_base_original = f"ESCRITO_{valor_columna_a}"
                nombre_base_original = re.sub(r'[<>:"/\\|?*]', '_', nombre_base_original).strip()

                # 2. Determinar el nombre final (base + sufijo) según si se firma o no
                if not self.sign_document.get():
                    # Sin firmar: ESCRITO_{valor}_no_firmado.pdf
                    nombre_final_base = f"{nombre_base_original}_no_firmado"
                else:
                    # Firmado: ESCRITO_{valor}.pdf
                    nombre_final_base = nombre_base_original
                
                # 3. Lógica de renombrado para encontrar un nombre de archivo único
                contador = 1
                nombre_base_para_usar = nombre_final_base
                ruta_final_pdf = os.path.join(output_dir, f"{nombre_base_para_usar}.pdf")

                while os.path.exists(ruta_final_pdf):
                    nombre_base_para_usar = f"{nombre_final_base}_{contador}"
                    ruta_final_pdf = os.path.join(output_dir, f"{nombre_base_para_usar}.pdf")
                    contador += 1

                # 4. Generar y firmar (si es necesario)
                if self.sign_document.get():
                    # Si se firma, se crea un PDF temporal que luego se elimina
                    ruta_pdf_temporal = os.path.join(output_dir, f"temp_{os.path.basename(ruta_final_pdf)}")
                    
                    # Generar el PDF inicial (con el sello visual) en la ruta temporal
                    self.generar_pdf(self.doc_path.get(), fila, ruta_pdf_temporal, cert_data=cert)
                    
                    # Firmar el PDF temporal y guardarlo en la ruta final definitiva
                    self.log("Firmando PDF...")
                    self.firmar_pdf(ruta_pdf_temporal, ruta_final_pdf, private_key, cert)
                    
                    # Eliminar el archivo temporal
                    os.remove(ruta_pdf_temporal)
                else:
                    # Si no se firma, se genera el PDF directamente en su ruta final
                    self.generar_pdf(self.doc_path.get(), fila, ruta_final_pdf)
                
                self.update_progress(i + 1)
                
            self.log("\n✅ ¡Proceso completado con éxito!")
            
            success_message = f"Se han generado {total_filas} PDFs en la carpeta '{output_dir}'."
            self.root.after(0, lambda: messagebox.showinfo("Éxito", success_message))
            
        except Exception as e:
            error_message = f"Ha ocurrido un error:\n\n{str(e)}"
            self.log(f"❌ Error: {str(e)}")
            
            self.root.after(0, lambda: messagebox.showerror("Error", error_message))
            
        finally:
            self.root.after(0, lambda: self.process_btn.config(state='normal'))
            self.update_progress(0)

    def extraer_campos_doc(self, doc_path):
        # Leer documento Word y extraer campos entre «»
        doc = Document(doc_path)
        campos = set()
        patron = re.compile(r'«(.*?)»')
        
        for paragraph in doc.paragraphs:
            coincidencias = patron.findall(paragraph.text)
            campos.update(coincidencias)
            
        # Buscar también en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    coincidencias = patron.findall(cell.text)
                    campos.update(coincidencias)
                    
        return list(campos)
        
    def validar_csv(self, csv_path, campos_requeridos):
        try:
            encoding_usada = 'utf-8-sig'
            try:
                # Intenta leer con la codificación estándar (UTF-8)
                with open(csv_path, mode='r', encoding=encoding_usada, newline='') as archivo:
                    lector = csv.DictReader(archivo, delimiter=';') 
                    
                    cabeceras = lector.fieldnames
                    # Quitamos espacios extra que puedan tener las cabeceras
                    lector.fieldnames = [name.strip() for name in cabeceras] 
                    datos = list(lector)

            except UnicodeDecodeError:
                encoding_usada = 'windows-1252'
                self.log(f"... Falló la lectura como UTF-8, reintentando con '{encoding_usada}'...")
                with open(csv_path, mode='r', encoding=encoding_usada, newline='') as archivo:
                    dialect = csv.Sniffer().sniff(archivo.read(1024), delimiters=',;')
                    archivo.seek(0)
                    lector = csv.DictReader(archivo, dialect=dialect)
                    cabeceras = lector.fieldnames
                    datos = list(lector)
            
            if not all(campo in cabeceras for campo in campos_requeridos):
                campos_faltantes = [c for c in campos_requeridos if c not in cabeceras]
                raise Exception(f"Faltan las siguientes columnas en el CSV: {', '.join(campos_faltantes)}")
            
            if not datos: 
                raise Exception("El archivo CSV está vacío.")
                
            return datos

        except Exception as e:
            raise Exception(f"Error al leer el CSV: {e}")

    def generar_pdf(self, doc_path, datos, output_pdf_path, cert_data=None):
        """
        Crea un DOCX temporal, lo rellena con datos, añade el sello si corresponde
        y lo convierte a PDF, eligiendo el método según el sistema operativo.
        """
        self.log("Creando documento personalizado...")
        doc = Document(doc_path)

        for p in doc.paragraphs:
            for key, value in datos.items():
                if f'«{key}»' in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if f'«{key}»' in inline[i].text:
                            text = inline[i].text.replace(f'«{key}»', str(value))
                            inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in datos.items():
                            if f'«{key}»' in inline[i].text:
                                text = inline[i].text.replace(f'«{key}»', str(value))
                                inline[i].text = text
        
        if cert_data:
            self.log("Añadiendo sello de firma al pie de página del DOCX...")
            self.add_signature_stamp_to_docx(doc, cert_data)

        temp_docx_path = output_pdf_path.replace('.pdf', '.docx')
        doc.save(temp_docx_path)
        
        self.log(f"Convirtiendo a PDF: {os.path.basename(output_pdf_path)}")
        
        try:
            sistema_operativo = platform.system()
            
            if sistema_operativo == "Windows":
                self.log("Usando docx2pdf (Windows)...")
                convert(temp_docx_path, output_pdf_path)
                
            elif sistema_operativo == "Linux":
                self.log("Usando LibreOffice (Linux)...")
                output_dir = os.path.dirname(output_pdf_path)
                subprocess.run(
                    ['lowriter', '--headless', '--convert-to', 'pdf', temp_docx_path, '--outdir', output_dir],
                    check=True, capture_output=True, text=True
                )
                generated_pdf_path = os.path.join(output_dir, os.path.basename(temp_docx_path).replace('.docx', '.pdf'))
                os.rename(generated_pdf_path, output_pdf_path)

            else:
                raise Exception(f"Sistema operativo no soportado: {sistema_operativo}")

        except FileNotFoundError:
            raise Exception("Comando no encontrado. En Windows, asegúrate de tener Word. En Linux, asegúrate de que LibreOffice esté instalado y 'lowriter' esté en el PATH.")
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error durante la conversión con LibreOffice: {e.stderr}")
        finally:
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)

    def add_signature_stamp_to_docx(self, doc, cert_data):
        try:
            signer_name = cert_data.subject.get_attributes_for_oid(x509.OID_COMMON_NAME)[0].value
        except (IndexError, AttributeError):
            signer_name = "Nombre Desconocido"

        signer_id = "ID Desconocido"
        try:
            match_id = re.search(r'\b([A-Z0-9]{9})\b', signer_name) # Patrón básico para NIF/CIF español
            if match_id:
                signer_id = match_id.group(1)
        except Exception:
            pass # No se pudo extraer, se queda en "ID Desconocido"
        
        fecha_actual_local = datetime.datetime.now().strftime("%Y.%m.%d %H:%M:%S")
        
        stamp_text = (
            f"Firmado digitalmente por\n"
            f"{signer_id} {signer_name}\n"
            f"Fecha: {fecha_actual_local}"
        )

        try:
            footer = doc.sections[0].footer
        except IndexError:
            section = doc.add_section()
            footer = section.footer

        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = paragraph.add_run(stamp_text)
        
        run.font.size = Pt(8) # Tamaño de fuente más pequeño
        run.font.name = 'Arial' # Tipo de fuente
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Alineado a la derecha
        #  borde superior al párrafo - separador visual
        # from docx.oxml.ns import qn
        # from docx.oxml import OxmlElement
        # pPr = paragraph._element.get_or_add_pPr()
        # pBdr = OxmlElement('w:pBdr')
        # pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDN', 'w:autoSpaceDE', 'w:wgBreakP', 'w:footnoteRef', 'w:endnoteRef', 'w:footnoteCont', 'w:endnoteCont', 'w:bare', 'w:noLnn', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:isLgl', 'w:cnfStyle', 'w:rPr', 'w:oMath')
        # top_border = OxmlElement('w:top')
        # top_border.set(qn('w:val'), 'single')
        # top_border.set(qn('w:sz'), '4') # Espesor del borde (1/8 de punto)
        # top_border.set(qn('w:space'), '1') # Espacio entre texto y borde
        # top_border.set(qn('w:color'), 'auto')
        # pBdr.append(top_border)
                
    def load_certificate(self, cert_path, password):
        try:
            with open(cert_path, 'rb') as f:
                p12_data = f.read()
            private_key, cert, _ = pkcs12.load_key_and_certificates(p12_data, password.encode() if password else None, default_backend())
            return private_key, cert
        except Exception as e:
            raise Exception(f"Error al cargar certificado. ¿Contraseña incorrecta? Detalle: {e}")
        
    def firmar_pdf(self, pdf_path, output_path, private_key, certificate):
        try:
            with open(pdf_path, 'rb') as f:
                pdf_data = f.read()
            
            # Si la versión de Python es 3.11 o superior...
            if sys.version_info >= (3, 11):
                fecha_utc = datetime.datetime.now(datetime.UTC)
            # Para versiones más antiguas
            else:
                from datetime import timezone
                fecha_utc = datetime.datetime.now(timezone.utc)
            
            fecha = fecha_utc.strftime("D:%Y%m%d%H%M%S+00'00'")
            
            dct = {
                'sigflags': 3, 
                'signingdate': fecha.encode(), 
            }
            
            signature = pdf.cms.sign(pdf_data, dct, private_key, certificate, [], 'sha256')
            
            with open(output_path, 'wb') as f:
                f.write(pdf_data + signature)
                
            self.log(f"📄 PDF firmado guardado como: {os.path.basename(output_path)}")
            
        except Exception as e:
            raise Exception(f"Error al firmar el PDF: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFGeneratorApp(root)
    root.mainloop()
